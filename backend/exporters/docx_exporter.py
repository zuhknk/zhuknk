"""Word 导出"""

import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_docx(data: dict) -> bytes:
    """将分析结果导出为 Word 文档"""
    doc = Document()

    # 标题
    title = doc.add_heading('App Review Analysis Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    reviews = data.get("reviews", [])
    findings = data.get("findings", [])
    requirements = data.get("requirements", [])
    test_cases = data.get("test_cases", [])

    # 概览
    doc.add_heading('Overview', level=1)
    table = doc.add_table(rows=5, cols=2, style='Light Shading Accent 1')
    cells = [
        ('Total Reviews', len(reviews)),
        ('Findings', len(findings)),
        ('Requirements', len(requirements)),
        ('Test Cases', len(test_cases)),
    ]
    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
    doc.add_paragraph()

    # 发现
    if findings:
        doc.add_heading('Findings', level=1)
        for f in findings:
            p = doc.add_paragraph()
            run = p.add_run(f"{f.get('topic', '')} [{f.get('severity', '')}]")
            run.bold = True
            doc.add_paragraph(f.get('description', ''))
            doc.add_paragraph()

    # 需求
    if requirements:
        doc.add_heading('Requirements', level=1)
        for r in requirements:
            p = doc.add_paragraph()
            run = p.add_run(f"{r.get('req_id', '')}: {r.get('title', '')} [{r.get('priority', '')}]")
            run.bold = True
            doc.add_paragraph(r.get('description', ''))
            if r.get('user_story'):
                p = doc.add_paragraph()
                run = p.add_run('User Story: ')
                run.italic = True
                p.add_run(r.get('user_story', ''))
            doc.add_paragraph()

    # 测试用例
    if test_cases:
        doc.add_heading('Test Cases', level=1)
        for tc in test_cases:
            p = doc.add_paragraph()
            run = p.add_run(f"{tc.get('case_id', '')}: {tc.get('title', '')}")
            run.bold = True
            steps = tc.get('steps', [])
            for i, step in enumerate(steps, 1):
                doc.add_paragraph(f'{i}. {step}', style='List Number')
            if tc.get('expected_result'):
                p = doc.add_paragraph()
                run = p.add_run('Expected: ')
                run.bold = True
                p.add_run(tc.get('expected_result', ''))
            doc.add_paragraph()

    # 保存到 bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
